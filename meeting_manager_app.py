# -*- coding: utf-8 -*-
"""
会议数据管理小程序（六字段分段模糊查询版，支持修改/删除，兼容居中会议名称）
"""
import sys
import re
import csv
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Tuple, Optional

from PyQt5 import QtCore, QtGui, QtWidgets
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Cm
import os, sys

if getattr(sys, 'frozen', False):  # 如果是打包后的 app
    bundle_dir = sys._MEIPASS if hasattr(sys, "_MEIPASS") else os.path.abspath(os.path.dirname(sys.executable))
    plugin_path = os.path.join(os.path.dirname(sys.executable), "../Plugins")
    os.environ["QT_PLUGIN_PATH"] = plugin_path




# DB_PATH = "meetings.db"

def external_resource(relative_path):
    """
    获取 .app 外部资源的路径
    """
    if getattr(sys, 'frozen', False):  # 如果是打包后的应用
        # __file__ 指向的是 Contents/MacOS/MeetingManager
        base_path = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(sys.executable))))
        # 这里的 base_path = dist/
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

DB_PATH = external_resource("meetings.db")

# ----------------------------- 数据库层 -----------------------------
class DBManager:
    def __init__(self, db_path: str = DB_PATH):
        self.conn = sqlite3.connect(db_path)
        self.conn.row_factory = sqlite3.Row
        self._init_schema()

    def _init_schema(self):
        cur = self.conn.cursor()
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS meetings (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                date TEXT NOT NULL,
                location TEXT NOT NULL,
                attendees TEXT NOT NULL,
                topic TEXT,
                content TEXT NOT NULL,
                raw_text TEXT NOT NULL,
                UNIQUE(title, date)
            );
            """
        )
        self.conn.commit()

    def insert_meeting(self, rec: Dict) -> Tuple[bool, Optional[str]]:
        try:
            cur = self.conn.cursor()
            cur.execute(
                """
                INSERT INTO meetings (title, date, location, attendees, topic, content, raw_text)
                VALUES (?, ?, ?, ?, ?, ?, ?);
                """,
                (
                    rec["title"],
                    rec["date"],
                    rec["location"],
                    rec["attendees"],
                    rec.get("topic", ""),
                    rec["content"],
                    rec["raw_text"],
                ),
            )
            self.conn.commit()
            return True, None
        except sqlite3.IntegrityError:
            return False, "重复记录（同“会议名称 + 会议时间”）已跳过。"
        except Exception as e:
            return False, f"插入失败：{e}"

    def update_meeting(self, rec_id: int, rec: Dict) -> Tuple[bool, Optional[str]]:
        try:
            cur = self.conn.cursor()
            cur.execute(
                """
                UPDATE meetings
                SET title=?, date=?, location=?, attendees=?, topic=?, content=?, raw_text=?
                WHERE id=?
                """,
                (
                    rec["title"], rec["date"], rec["location"], rec["attendees"],
                    rec.get("topic", ""), rec["content"], rec["raw_text"], rec_id
                ),
            )
            self.conn.commit()
            return True, None
        except sqlite3.IntegrityError:
            return False, "修改后与已有记录重复（同“会议名称+会议时间”）。"
        except Exception as e:
            return False, f"更新失败：{e}"

    def delete_meeting(self, rec_id: int) -> Tuple[bool, Optional[str]]:
        try:
            cur = self.conn.cursor()
            cur.execute("DELETE FROM meetings WHERE id=?", (rec_id,))
            self.conn.commit()
            return True, None
        except Exception as e:
            return False, f"删除失败：{e}"

    def get_id_by_title_date(self, title: str, date: str) -> Optional[int]:
        cur = self.conn.cursor()
        cur.execute("SELECT id FROM meetings WHERE title=? AND date=?", (title, date))
        row = cur.fetchone()
        return row["id"] if row else None


    def segmented_search_exact_six(
        self,
        title_kw: str = "",
        date_kw: str = "",
        location_kw: str = "",
        attendees_kw: str = "",
        topic_kw: str = "",
        content_kw: str = "",
    ) -> List[sqlite3.Row]:
        clauses = []
        params = []

        def like_clause(field: str, kw: str):
            clauses.append(f"{field} LIKE ?")
            params.append(f"%{kw.strip()}%")

        if title_kw.strip():
            like_clause("title", title_kw)
        if location_kw.strip():
            like_clause("location", location_kw)
        if attendees_kw.strip():
            like_clause("attendees", attendees_kw)
        if topic_kw.strip():
            like_clause("topic", topic_kw)
        if content_kw.strip():
            like_clause("content", content_kw)

        if date_kw.strip():
            prefix = normalize_date_prefix(date_kw.strip())
            if prefix:
                clauses.append("date LIKE ?")
                params.append(prefix + "%")

        where_sql = " AND ".join(clauses) if clauses else "1=1"
        sql = f"""
            SELECT * FROM meetings
            WHERE {where_sql}
            ORDER BY date DESC, id DESC
        """
        cur = self.conn.cursor()
        cur.execute(sql, tuple(params))
        return cur.fetchall()

    def close(self):
        self.conn.close()


# ----------------------------- 日期工具 -----------------------------
DATE_RE_CN_FULL = re.compile(r"^\s*(?P<y>\d{4})年(?P<m>\d{1,2})月(?P<d>\d{1,2})日\s*$")
DATE_RE_STD_FULL = re.compile(r"^\s*(?P<y>\d{4})-(?P<m>\d{1,2})-(?P<d>\d{1,2})\s*$")
# 年月日后面允许跟任意内容（比如时间、文字）
DATE_RE_CN_PREFIX = re.compile(r"^\s*(?P<y>\d{4})年(?P<m>\d{1,2})月(?P<d>\d{1,2})日")
DATE_RE_STD_PREFIX = re.compile(r"^\s*(?P<y>\d{4})-(?P<m>\d{1,2})-(?P<d>\d{1,2})")


def pad2(n: int) -> str:
    return f"{n:02d}"


def normalize_datetime(date_str: str) -> Optional[str]:
    """
    解析会议时间，统一存为 '2025年10月15日' 这种格式，
    如果有后缀，就保留（例如：'2025年10月15日 下午3点'）。
    """
    s = date_str.strip()

    # 中文日期 + 后缀
    m = DATE_RE_CN_PREFIX.match(s)
    if m:
        y, mm, d = int(m.group("y")), int(m.group("m")), int(m.group("d"))
        try:
            date_part = f"{y}年{mm}月{d}日"
            suffix = s[m.end():].strip()
            return f"{date_part} {suffix}" if suffix else date_part
        except ValueError:
            return None

    # 标准日期 + 后缀
    m = DATE_RE_STD_PREFIX.match(s)
    if m:
        y, mm, d = int(m.group("y")), int(m.group("m")), int(m.group("d"))
        try:
            date_part = f"{y}年{mm}月{d}日"
            suffix = s[m.end():].strip()
            return f"{date_part} {suffix}" if suffix else date_part
        except ValueError:
            return None

    return None

# def normalize_datetime(date_str: str) -> Optional[str]:
#     """
#     解析会议时间：提取前面的日期部分，后面有内容就原样拼上
#     输入示例: "2022年1月4日10:00" -> "2022-01-04 10:00"
#              "2022-01-04 下午3点"   -> "2022-01-04 下午3点"
#     """
#     s = date_str.strip()

#     # 中文日期 + 后缀
#     # m = DATE_RE_CN_FULL.match(s)
#     m = DATE_RE_CN_PREFIX.match(s)
#     if m:
#         y, mm, d = int(m.group("y")), int(m.group("m")), int(m.group("d"))
#         try:
#             date_part = datetime(y, mm, d).strftime("%Y-%m-%d")
#             suffix = s[m.end():].strip()  # 拿掉日期后剩下的
#             return f"{date_part} {suffix}" if suffix else date_part
#         except ValueError:
#             return None

#     # 标准日期 + 后缀
#     # m = DATE_RE_STD_FULL.match(s)
#     m = DATE_RE_STD_PREFIX.match(s)
#     if m:
#         y, mm, d = int(m.group("y")), int(m.group("m")), int(m.group("d"))
#         try:
#             date_part = datetime(y, mm, d).strftime("%Y-%m-%d")
#             suffix = s[m.end():].strip()
#             return f"{date_part} {suffix}" if suffix else date_part
#         except ValueError:
#             return None

#     return None


# def normalize_date_prefix(date_in: str) -> Optional[str]:
#     s = date_in.strip()
#     m = DATE_RE_CN_FULL.match(s)
#     if m:
#         try:
#             return datetime(int(m.group("y")), int(m.group("m")), int(m.group("d"))).strftime("%Y-%m-%d")
#         except ValueError:
#             return None
#     return None
def normalize_date_prefix(date_in: str) -> Optional[str]:
    s = date_in.strip()

    # 中文日期开头
    m = DATE_RE_CN_FULL.match(s)
    if m:
        try:
            return datetime(
                int(m.group("y")), int(m.group("m")), int(m.group("d"))
            ).strftime("%Y-%m-%d")
        except ValueError:
            return None

    # 标准日期开头
    m = DATE_RE_STD_FULL.match(s)
    if m:
        try:
            return datetime(
                int(m.group("y")), int(m.group("m")), int(m.group("d"))
            ).strftime("%Y-%m-%d")
        except ValueError:
            return None

    return None



# ----------------------------- 解析 Word -----------------------------
def parse_docx_to_records(path: Path) -> List[Tuple[Optional[Dict], Optional[str], Optional[str]]]:
    """
    一个 Word 文件可能有多个会议，用分页符区分。
    返回多个 (record, error, raw_text)
    """
    all_meetings = []
    try:
        meetings = split_docx_by_page(path)
    except Exception as e:
        return [(None, f"无法解析Word：{e}", None)]

    for lines in meetings:
        raw_text = "\n".join(lines)
        if not lines:
            all_meetings.append((None, "文档为空", raw_text))
            continue

        title = lines[0]
        date_str = location = attendees = topic = ""
        content_lines = []
        content_started = False

        def safe_extract(ln: str) -> str:
            parts = ln.split("：", 1)
            return parts[1].strip() if len(parts) > 1 else ""

        for ln in lines[1:]:
            if ln.startswith("会议时间"):
                date_str = safe_extract(ln)
            elif ln.startswith("会议地点"):
                location = safe_extract(ln)
            elif ln.startswith("参会人员"):
                attendees = safe_extract(ln)
            elif ln.startswith("会议议题"):
                topic = safe_extract(ln)
            elif ln.startswith("会议内容"):
                content_started = True
            else:
                if content_started:
                    content_lines.append(ln)

        if not date_str or not location or not attendees:
            all_meetings.append((None, "缺少必要字段（会议时间/地点/人员）", raw_text))
            continue

        date_norm = normalize_datetime(date_str)
        if not date_norm:
            all_meetings.append((None, "会议时间格式不正确", raw_text))
            continue

        content = "\n".join(content_lines).strip()
        rec = {
            "title": title,
            "date": date_norm,
            "location": location,
            "attendees": attendees,
            "topic": topic,
            "content": content,
            "raw_text": raw_text,
        }
        all_meetings.append((rec, None, None))

    return all_meetings

# def parse_docx_to_record(path: Path) -> Tuple[Optional[Dict], Optional[str], Optional[str]]:
#     try:
#         doc = DocxDocument(str(path))
#     except Exception as e:
#         return None, f"无法解析Word：{e}", None

#     lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
#     raw_text = "\n".join(lines)
#     if not lines:
#         return None, "文档为空", raw_text

#     # 第一行作为会议名称
#     title = lines[0]

#     date_str = location = attendees = topic = ""
#     content_lines = []
#     content_started = False

#     def safe_extract(ln: str) -> str:
#         parts = ln.split("：", 1)
#         return parts[1].strip() if len(parts) > 1 else ""

#     for ln in lines[1:]:
#         if ln.startswith("会议时间"):
#             date_str = safe_extract(ln)
#         elif ln.startswith("会议地点"):
#             location = safe_extract(ln)
#         elif ln.startswith("参会人员"):
#             attendees = safe_extract(ln)
#         elif ln.startswith("会议议题"):
#             topic = safe_extract(ln)
#         elif ln.startswith("会议内容"):
#             content_started = True
#         else:
#             if content_started:
#                 content_lines.append(ln)

#     if not date_str or not location or not attendees:
#         return None, "缺少必要字段（会议时间/地点/人员）", raw_text

#     # date_norm = normalize_date(date_str)
#     date_norm = normalize_datetime(date_str)
#     if not date_norm:
#         return None, "会议时间格式不正确", raw_text

#     content = "\n".join(content_lines).strip()

#     return {
#         "title": title,
#         "date": date_norm,
#         "location": location,
#         "attendees": attendees,
#         "topic": topic,
#         "content": content,
#         "raw_text": raw_text,
#     }, None, None

def split_docx_by_page(path: Path) -> List[List[str]]:
    """
    读取 Word 文档，并按照手动分页符 (Ctrl+Enter) 分割成多个会议。
    每个会议是一组文本行 (list[str])。
    """
    from docx.oxml.ns import qn

    doc = DocxDocument(str(path))
    meetings, current_lines = [], []

    for para in doc.paragraphs:
        text = para.text.strip()
        has_page_break = False

        for run in para.runs:
            for br in run._element.findall(".//w:br", namespaces=run._element.nsmap):
                if br.get(qn("w:type")) == "page":
                    has_page_break = True

        if text:
            current_lines.append(text)

        if has_page_break:  # 遇到分页符 → 结束当前会议
            if current_lines:
                meetings.append(current_lines)
                current_lines = []

    if current_lines:
        meetings.append(current_lines)

    return meetings



# ----------------------------- GUI 层 -----------------------------
class ManualEntryDialog(QtWidgets.QDialog):
    def __init__(self, parent=None, db: DBManager = None, rec_id: Optional[int] = None):
        super().__init__(parent)
        self.setWindowTitle("手动新增/修改会议")
        self.setMinimumWidth(560)
        self.db = db
        self.rec_id = rec_id  # 如果是修改，带上 id

        form = QtWidgets.QFormLayout()

        self.ed_title = QtWidgets.QLineEdit()
        self.ed_date = QtWidgets.QLineEdit()
        self.ed_location = QtWidgets.QLineEdit()
        self.ed_attendees = QtWidgets.QLineEdit()
        self.ed_topic = QtWidgets.QLineEdit()
        self.ed_content = QtWidgets.QTextEdit()
        self.ed_content.setMinimumHeight(140)

        form.addRow("会议名称：", self.ed_title)
        form.addRow("会议时间：", self.ed_date)
        form.addRow("会议地点：", self.ed_location)
        form.addRow("参会人员：", self.ed_attendees)
        form.addRow("会议议题：", self.ed_topic)
        # form.addRow("会议内容：", self.ed_content)
        form.addRow(QtWidgets.QLabel("会议内容："))
        form.addRow(self.ed_content)

        btns = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Ok | QtWidgets.QDialogButtonBox.Cancel)
        btns.button(QtWidgets.QDialogButtonBox.Ok).setText("保存")
        btns.accepted.connect(self.on_save_clicked)  # 修改：绑定保存函数
        btns.rejected.connect(self.reject)  # Cancel 还是关闭

        layout = QtWidgets.QVBoxLayout(self)
        layout.addLayout(form)
        layout.addWidget(btns)

        for w in self.findChildren(QtWidgets.QLineEdit):
            w.setFixedHeight(36)

    def set_record(self, rec: sqlite3.Row):
        self.ed_title.setText(rec["title"])
        self.ed_date.setText(rec["date"])
        self.ed_location.setText(rec["location"])
        self.ed_attendees.setText(rec["attendees"])
        self.ed_topic.setText(rec["topic"])
        self.ed_content.setPlainText(rec["content"])

    def collect_record(self) -> Optional[Dict]:
        title = self.ed_title.text().strip()
        date_in = self.ed_date.text().strip()
        location = self.ed_location.text().strip()
        attendees = self.ed_attendees.text().strip()
        topic = self.ed_topic.text().strip()
        content = self.ed_content.toPlainText().strip()

        # 必填：会议名称、会议时间
        if not title or not date_in:
            QtWidgets.QMessageBox.warning(self, "提示", "会议名称和会议时间不能为空。")
            return None

        date_norm = normalize_datetime(date_in)
        if not date_norm:
            QtWidgets.QMessageBox.warning(
                self,
                "提示",
                "会议时间格式不正确（例如：2022年1月4日10:00 或 2022-01-04 10:00）。"
            )
            return None

        raw_text = (
            f"{title}\n"
            f"会议时间：{date_norm}\n"
            f"会议地点：{location}\n"
            f"参会人员：{attendees}\n"
            f"会议议题：{topic}\n"
            f"会议内容：\n{content}"
        )

        return {
            "title": title,
            "date": date_norm,
            "location": location,
            "attendees": attendees,
            "topic": topic,
            "content": content,
            "raw_text": raw_text,
        }

    # def on_save_clicked(self):
    #     rec = self.collect_record()
    #     if not rec:
    #         return
    #     if self.db:
    #         if self.rec_id is None:  # 新增
    #             ok, msg = self.db.insert_meeting(rec)
    #         else:  # 修改
    #             ok, msg = self.db.update_meeting(self.rec_id, rec)

    #         if ok:
    #             QtWidgets.QMessageBox.information(self, "成功", "已保存。")
    #         else:
    #             QtWidgets.QMessageBox.warning(self, "提示", msg or "保存失败")

    def on_save_clicked(self):
        rec = self.collect_record()
        if not rec:
            return
        if not self.db:
            return

        if self.rec_id is None:  # 第一次保存：插入
            ok, msg = self.db.insert_meeting(rec)
            if ok:
                # 关键：把刚插入的记录 id 取回来，切换为“修改模式”
                rid = self.db.get_id_by_title_date(rec["title"], rec["date"])
                if rid is not None:
                    self.rec_id = rid  # 之后再点保存就会走 update
                QtWidgets.QMessageBox.information(self, "成功", "已保存。")
                # 可选：让主界面刷新列表（若父窗口有该方法）
                parent = self.parent()
                if parent and hasattr(parent, "refresh_results"):
                    parent.refresh_results()
            else:
                QtWidgets.QMessageBox.warning(self, "提示", msg or "保存失败")
        else:  # 后续保存：更新
            ok, msg = self.db.update_meeting(self.rec_id, rec)
            if ok:
                QtWidgets.QMessageBox.information(self, "成功", "已保存修改。")
                parent = self.parent()
                if parent and hasattr(parent, "refresh_results"):
                    parent.refresh_results()
            else:
                QtWidgets.QMessageBox.warning(self, "提示", msg or "保存失败")




class DetailsDialog(QtWidgets.QDialog):
    def __init__(self, rec: sqlite3.Row, parent=None):
        super().__init__(parent)
        self.setWindowTitle("会议详情")
        self.setMinimumWidth(640)

        text = (
            f"{rec['title']}\n"
            f"会议时间：{rec['date']}\n"
            f"会议地点：{rec['location']}\n"
            f"参会人员：{rec['attendees']}\n"
            f"会议议题：{rec['topic']}\n"
            f"会议内容：\n{rec['content']}"
        )
        te = QtWidgets.QTextEdit()
        te.setReadOnly(True)
        te.setPlainText(text)

        layout = QtWidgets.QVBoxLayout(self)
        layout.addWidget(te)


class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("会议数据管理小程序")
        self.resize(980, 660)

        self.db = DBManager()
        self.last_results: List[sqlite3.Row] = []

        # 顶部按钮
        btn_import = QtWidgets.QPushButton("导入Word")
        btn_add = QtWidgets.QPushButton("手动新增")
        btn_export = QtWidgets.QPushButton("导出")

        btn_import.clicked.connect(self.on_import_clicked)
        btn_add.clicked.connect(self.on_add_clicked)
        btn_export.clicked.connect(self.on_export_clicked)

        top_bar = QtWidgets.QHBoxLayout()
        top_bar.addWidget(btn_import)
        top_bar.addWidget(btn_add)
        top_bar.addWidget(btn_export)
        top_bar.addStretch()

        # 字体调节区
        btn_dec_font = QtWidgets.QPushButton("-")
        btn_inc_font = QtWidgets.QPushButton("+")
        self.font_size_box = QtWidgets.QLineEdit("20")
        self.font_size_box.setFixedWidth(50)
        self.font_size_box.setAlignment(QtCore.Qt.AlignCenter)

        def apply_font_size(size: int):
            font = QtGui.QFont("SimSun", size)
            QtWidgets.QApplication.instance().setFont(font)
            self.font_size_box.setText(str(size))

        btn_dec_font.clicked.connect(lambda: apply_font_size(max(8, int(self.font_size_box.text()) - 1)))
        btn_inc_font.clicked.connect(lambda: apply_font_size(int(self.font_size_box.text()) + 1))
        self.font_size_box.editingFinished.connect(
            lambda: apply_font_size(max(8, int(self.font_size_box.text())))
        )

        top_bar.addWidget(btn_dec_font)
        top_bar.addWidget(self.font_size_box)
        top_bar.addWidget(btn_inc_font)








        # 中部：六字段查询
        grid = QtWidgets.QGridLayout()
        self.q_title = QtWidgets.QLineEdit()
        self.q_date = QtWidgets.QLineEdit()
        self.q_location = QtWidgets.QLineEdit()
        self.q_attendees = QtWidgets.QLineEdit()
        self.q_topic = QtWidgets.QLineEdit()
        self.q_content = QtWidgets.QLineEdit()

        grid.addWidget(QtWidgets.QLabel("会议名称"), 0, 0); grid.addWidget(self.q_title, 0, 1, 1, 3)
        grid.addWidget(QtWidgets.QLabel("会议时间"), 1, 0); grid.addWidget(self.q_date, 1, 1, 1, 3)
        grid.addWidget(QtWidgets.QLabel("会议地点"), 2, 0); grid.addWidget(self.q_location, 2, 1, 1, 3)
        grid.addWidget(QtWidgets.QLabel("参会人员"), 3, 0); grid.addWidget(self.q_attendees, 3, 1, 1, 3)
        grid.addWidget(QtWidgets.QLabel("会议议题"), 4, 0); grid.addWidget(self.q_topic, 4, 1, 1, 3)
        grid.addWidget(QtWidgets.QLabel("会议内容"), 5, 0); grid.addWidget(self.q_content, 5, 1, 1, 3)

        btn_search = QtWidgets.QPushButton("查询")
        btn_search.clicked.connect(self.on_search_clicked)
        grid.addWidget(btn_search, 0, 4, 1, 1)

        # 底部：结果列表
        self.table = QtWidgets.QTableWidget(0, 3)
        self.table.setHorizontalHeaderLabels(["会议名称", "时间", "会议议题"])
        self.table.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        self.table.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection)
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.doubleClicked.connect(self.on_row_double_clicked)

        # 右键菜单
        self.table.setContextMenuPolicy(QtCore.Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.on_table_context_menu)
        # self.table.setSortingEnabled(True)


        central = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout(central)
        layout.addLayout(top_bar)
        layout.addLayout(grid)
        layout.addWidget(self.table)
        self.setCentralWidget(central)

        for w in self.findChildren(QtWidgets.QLineEdit):
            w.setFixedHeight(32)

        # 初始展示
        self.refresh_results()

    def refresh_results(self):
        self.last_results = self.db.segmented_search_exact_six(
            title_kw=self.q_title.text(), date_kw=self.q_date.text(),
            location_kw=self.q_location.text(), attendees_kw=self.q_attendees.text(),
            topic_kw=self.q_topic.text(), content_kw=self.q_content.text(),
        )
        self.table.setRowCount(len(self.last_results))
        for r, row in enumerate(self.last_results):
            self.table.setItem(r, 0, QtWidgets.QTableWidgetItem(row["title"]))
            self.table.setItem(r, 1, QtWidgets.QTableWidgetItem(row["date"]))
            self.table.setItem(r, 2, QtWidgets.QTableWidgetItem(row["topic"]))

    def on_search_clicked(self):
        self.refresh_results()

    def on_row_double_clicked(self):
        r = self.table.currentRow()
        if r < 0 or r >= len(self.last_results):
            return
        dlg = DetailsDialog(self.last_results[r], self)
        dlg.exec_()

    def on_table_context_menu(self, pos):
        row = self.table.currentRow()
        if row < 0 or row >= len(self.last_results):
            return
        rec = self.last_results[row]
        menu = QtWidgets.QMenu(self)
        act_edit = menu.addAction("修改")
        act_del = menu.addAction("删除")
        action = menu.exec_(self.table.viewport().mapToGlobal(pos))
        if action == act_edit:
            self.edit_record(rec)
        elif action == act_del:
            self.delete_record(rec)

    # def edit_record(self, rec):
    #     dlg = ManualEntryDialog(self)
    #     dlg.set_record(rec)
    #     if dlg.exec_() == QtWidgets.QDialog.Accepted:
    #         new_rec = dlg.get_record()
    #         if not new_rec:
    #             return
    #         ok, msg = self.db.update_meeting(rec["id"], new_rec)
    #         if ok:
    #             QtWidgets.QMessageBox.information(self, "成功", "已修改。")
    #             self.refresh_results()
    #         else:
    #             QtWidgets.QMessageBox.warning(self, "提示", msg or "修改失败")

    def edit_record(self, rec):
        dlg = ManualEntryDialog(self, db=self.db, rec_id=rec["id"])
        dlg.set_record(rec)
        dlg.exec_()
        self.refresh_results()

    def delete_record(self, rec):
        reply = QtWidgets.QMessageBox.question(
            self, "确认删除", f"确定要删除会议《{rec['title']}》吗？",
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            ok, msg = self.db.delete_meeting(rec["id"])
            if ok:
                QtWidgets.QMessageBox.information(self, "成功", "已删除。")
                self.refresh_results()
            else:
                QtWidgets.QMessageBox.warning(self, "提示", msg or "删除失败")

    # def on_add_clicked(self):
    #     dlg = ManualEntryDialog(self)
    #     if dlg.exec_() == QtWidgets.QDialog.Accepted:
    #         rec = dlg.get_record()
    #         if not rec:
    #             return
    #         ok, msg = self.db.insert_meeting(rec)
    #         if ok:
    #             QtWidgets.QMessageBox.information(self, "成功", "已新增。")
    #             self.refresh_results()
    #         else:
    #             QtWidgets.QMessageBox.warning(self, "提示", msg or "新增失败")
    def on_add_clicked(self):
        dlg = ManualEntryDialog(self, db=self.db)
        dlg.exec_()
        self.refresh_results()

    def on_import_clicked(self):
        files, _ = QtWidgets.QFileDialog.getOpenFileNames(
            self, "选择要导入的 Word 文件", "", "Word 文档 (*.docx)"
        )
        if not files:
            return

        errors, ok_count, dup_count = [], 0, 0
        for f in files:
            records = parse_docx_to_records(Path(f))  # ⚠️ 注意这里换成 parse_docx_to_records
            for rec, err, raw in records:
                if err:
                    errors.append((Path(f).name, err, raw or ""))
                    continue
                ok, msg = self.db.insert_meeting(rec)
                if ok:
                    ok_count += 1
                elif msg and "重复记录" in msg:
                    dup_count += 1
                else:
                    errors.append((Path(f).name, msg or "未知错误", rec.get("raw_text", "")))

        msg_parts = [f"成功导入：{ok_count} 条"]
        if dup_count:
            msg_parts.append(f"重复跳过：{dup_count} 条")
        if errors:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            err_csv = Path.cwd() / f"import_errors_{ts}.csv"
            with open(err_csv, "w", newline="", encoding="utf-8") as fp:
                w = csv.writer(fp)
                w.writerow(["文件名", "错误原因", "原始文本"])
                w.writerows(errors)
            msg_parts.append(f"错误明细已保存：{err_csv}")

        QtWidgets.QMessageBox.information(self, "导入完成", "\n".join(msg_parts))
        self.refresh_results()

    def on_export_clicked(self):
        if not self.last_results:
            QtWidgets.QMessageBox.information(self, "提示", "没有可导出的查询结果。")
            return
        ts = datetime.now().strftime("%Y%m%d")
        default_name = f"导出_{ts}.docx"
        out_path, _ = QtWidgets.QFileDialog.getSaveFileName(
            self, "保存导出文件", default_name, "Word 文档 (*.docx)"
        )
        if not out_path:
            return

        doc = DocxDocument()

        # 设置全局默认字体：SimSun + 14号
        style = doc.styles['Normal']
        style.font.name = 'FangSong'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
        style.font.size = Pt(16)

        # def set_run_style(run, bold=False):
        #     run.font.name = "FangSong"
        #     run._element.rPr.rFonts.set(qn('w:eastAsia'), "FangSong")
        #     run.font.size = Pt(16)
        #     run.bold = bold
        #     return run
        def set_run_style(run, bold=False):
            if bold:
                # 黑体 16
                run.font.name = "SimHei"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "SimHei")
                run.font.size = Pt(16)
                run.bold = False   # 不要再加粗
            else:
                # 仿宋 16
                run.font.name = "FangSong"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "FangSong")
                run.font.size = Pt(16)
                run.bold = False
            return run


        # 工具函数：生成无段间距段落
        def add_para_no_space(text="", bold=False, indent=False, align=None):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            if indent:
                para.paragraph_format.first_line_indent = Cm(1)  # 首行缩进两个字
            if align is not None:
                para.alignment = align
            run = set_run_style(para.add_run(text), bold=bold)
            return para

        for idx, rec in enumerate(self.last_results):
            # 会议名称 居中 + 加粗
            # add_para_no_space(rec["title"], bold=True, align=1)
            # doc.add_paragraph()
            para = add_para_no_space(rec["title"], bold=True, align=1)
            para.paragraph_format.space_after = Pt(12)   # 调节间距大小




            # 会议时间
            para = add_para_no_space()
            set_run_style(para.add_run("会议时间："), bold=True)
            set_run_style(para.add_run(rec["date"] or ""))

            # 会议地点
            para = add_para_no_space()
            set_run_style(para.add_run("会议地点："), bold=True)
            set_run_style(para.add_run(rec["location"] or ""))

            # 参会人员
            para = add_para_no_space()
            set_run_style(para.add_run("参会人员："), bold=True)
            set_run_style(para.add_run(rec["attendees"] or ""))

            # 会议议题 整行加粗
            add_para_no_space("会议议题：" + (rec['topic'] or ""), bold=True)

            # 会议内容
            add_para_no_space("会议内容：", bold=True)

            for line in (rec["content"] or "").splitlines():
                add_para_no_space(line, indent=True)

            # 多个会议之间加分页符
            if idx != len(self.last_results) - 1:
                doc.add_page_break()

        doc.save(out_path)
        QtWidgets.QMessageBox.information(self, "成功", f"导出完成：{out_path}")

    # def on_export_clicked(self):
    #     if not self.last_results:
    #         QtWidgets.QMessageBox.information(self, "提示", "没有可导出的查询结果。")
    #         return
    #     ts = datetime.now().strftime("%Y%m%d")
    #     default_name = f"导出_{ts}.docx"
    #     out_path, _ = QtWidgets.QFileDialog.getSaveFileName(
    #         self, "保存导出文件", default_name, "Word 文档 (*.docx)"
    #     )
    #     if not out_path:
    #         return

    #     doc = DocxDocument()

    #     # 设置全局默认字体：SimSun + 14号
    #     style = doc.styles['Normal']
    #     style.font.name = 'SimSun'
    #     style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    #     style.font.size = Pt(14)

    #     def set_run_style(run, bold=False):
    #         run.font.name = "SimSun"
    #         run._element.rPr.rFonts.set(qn('w:eastAsia'), "SimSun")
    #         run.font.size = Pt(14)
    #         run.bold = bold
    #         return run

    #     for idx, rec in enumerate(self.last_results):
    #         # 会议名称 居中 + 加粗
    #         para = doc.add_paragraph()
    #         run = set_run_style(para.add_run(rec["title"]), bold=True)
    #         para.alignment = 1  # 居中

    #         # 会议时间
    #         para = doc.add_paragraph()
    #         set_run_style(para.add_run("会议时间："), bold=True)
    #         set_run_style(para.add_run(rec["date"] or ""))

    #         # 会议地点
    #         para = doc.add_paragraph()
    #         set_run_style(para.add_run("会议地点："), bold=True)
    #         set_run_style(para.add_run(rec["location"] or ""))

    #         # 参会人员
    #         para = doc.add_paragraph()
    #         set_run_style(para.add_run("参会人员："), bold=True)
    #         set_run_style(para.add_run(rec["attendees"] or ""))

    #         # 会议议题 整行加粗
    #         para = doc.add_paragraph()
    #         set_run_style(para.add_run("会议议题：" + (rec['topic'] or "")), bold=True)

    #         # 会议内容
    #         para = doc.add_paragraph()
    #         set_run_style(para.add_run("会议内容："), bold=True)

    #         # for line in (rec["content"] or "").splitlines():
    #         #     para = doc.add_paragraph()
    #         #     set_run_style(para.add_run(line))
    #         for line in (rec["content"] or "").splitlines():
    #             para = doc.add_paragraph()
    #             para.paragraph_format.first_line_indent = Cm(1)  # 首行缩进两个字（约 1cm）
    #             set_run_style(para.add_run(line))

    #         if idx != len(self.last_results) - 1:
    #             doc.add_page_break()

    #     doc.save(out_path)
    #     QtWidgets.QMessageBox.information(self, "成功", f"导出完成：{out_path}")


def main():
    QtWidgets.QApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling, True)
    app = QtWidgets.QApplication(sys.argv)
    font = QtGui.QFont("SimSun", 20)
    app.setFont(font)

    win = MainWindow()
    win.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
